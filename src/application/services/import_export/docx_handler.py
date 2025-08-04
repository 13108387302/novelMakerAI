#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX格式处理器

处理Microsoft Word文档格式的项目和文档导入导出
"""

from pathlib import Path
from typing import List, Optional
from datetime import datetime

from .base import BaseFormatHandler, ExportOptions, ImportOptions
from src.domain.entities.project import Project
from src.domain.entities.document import Document, DocumentType
from src.shared.utils.logger import get_logger

logger = get_logger(__name__)

# 检查DOCX库是否可用
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx库不可用，DOCX格式处理器将无法正常工作")


class DocxFormatHandler(BaseFormatHandler):
    """
    DOCX格式处理器

    处理Microsoft Word文档格式的项目和文档导入导出操作。
    使用python-docx库生成专业的Word文档。

    实现方式：
    - 使用python-docx库处理Word文档
    - 支持丰富的文本格式和样式
    - 提供章节分页和目录结构
    - 支持表格和统计信息展示
    - 处理文档元数据和属性

    Features:
        - 项目信息格式化输出
        - 文档内容分章节组织
        - 统计信息表格展示
        - 支持样式和格式设置
        - 文档导入和解析功能
    """

    def get_supported_extensions(self) -> List[str]:
        """
        获取支持的文件扩展名

        Returns:
            List[str]: 支持的文件扩展名列表
        """
        return ['.docx']

    def get_format_name(self) -> str:
        """
        获取格式名称

        Returns:
            str: 格式的显示名称
        """
        return "Microsoft Word文档"

    async def _do_export_project(self, project: Project, output_path: Path, options: ExportOptions) -> bool:
        """导出项目为DOCX格式"""
        if not DOCX_AVAILABLE:
            logger.error("DOCX库不可用，无法导出DOCX格式")
            return False
            
        try:
            # 创建Word文档
            doc = DocxDocument()
            
            # 设置文档标题
            title = doc.add_heading(project.name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加项目描述
            if hasattr(project, 'description') and project.description:
                doc.add_heading('项目描述', level=1)
                doc.add_paragraph(project.description)
                doc.add_page_break()
            
            # 添加导出信息
            if options.include_metadata:
                doc.add_heading('导出信息', level=1)
                info_para = doc.add_paragraph()
                info_para.add_run(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                info_para.add_run(f"导出工具: AI小说编辑器\n")
                info_para.add_run(f"格式版本: DOCX")
                doc.add_page_break()
            
            # 获取并添加文档内容
            if options.include_documents:
                try:
                    documents = await self.service.document_repository.get_by_project_id(project.id)
                    
                    if documents:
                        doc.add_heading('文档内容', level=1)
                        
                        for i, document in enumerate(documents, 1):
                            # 添加文档标题
                            doc.add_heading(f"{i}. {document.title}", level=2)
                            
                            # 添加文档内容
                            content_paragraphs = document.content.split('\n\n')
                            for paragraph_text in content_paragraphs:
                                if paragraph_text.strip():
                                    doc.add_paragraph(paragraph_text.strip())
                            
                            # 添加分页符（除了最后一个文档）
                            if i < len(documents):
                                doc.add_page_break()
                                
                except Exception as e:
                    logger.warning(f"获取项目文档失败: {e}")
                    doc.add_paragraph("无法获取文档内容")
            
            # 添加统计信息
            if options.include_statistics:
                doc.add_page_break()
                doc.add_heading('统计信息', level=1)
                
                # 计算统计信息
                total_chars = 0
                total_words = 0
                
                try:
                    documents = await self.service.document_repository.get_by_project_id(project.id)
                    for document in documents:
                        total_chars += len(document.content)
                        total_words += len(document.content.split())
                except:
                    pass
                
                stats_para = doc.add_paragraph()
                stats_para.add_run(f"总字符数: {total_chars:,}\n")
                stats_para.add_run(f"总词数: {total_words:,}\n")
                stats_para.add_run(f"文档数量: {len(documents) if 'documents' in locals() else 0}")
            
            # 保存文档
            doc.save(str(output_path))
            
            logger.info(f"项目已导出为DOCX: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"导出DOCX失败: {e}")
            return False

    async def _do_import_project(self, input_path: Path, options: ImportOptions) -> Optional[Project]:
        """从DOCX导入项目"""
        if not DOCX_AVAILABLE:
            logger.error("DOCX库不可用，无法导入DOCX格式")
            return None
            
        try:
            # 验证文件扩展名
            if not self._validate_file_extension(input_path, self.get_supported_extensions()):
                logger.error(f"不支持的文件格式: {input_path.suffix}")
                return None

            # 读取DOCX文件
            doc = DocxDocument(str(input_path))
            
            # 提取项目名称（使用文件名或第一个标题）
            project_name = input_path.stem
            
            # 尝试从文档中提取标题
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading') and paragraph.text.strip():
                    project_name = paragraph.text.strip()
                    break
            
            # 提取所有文本内容
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            content = '\n\n'.join(full_text)
            
            # 创建项目
            project = Project(
                id="",  # 将由仓储分配
                name=project_name,
                description=f"从DOCX文件导入: {input_path.name}"
            )
            
            # 创建文档包含内容
            document = Document(
                id="",  # 将由仓储分配
                title=f"{project_name} - 内容",
                content=content,
                type=DocumentType.CHAPTER
            )
            
            # 保存文档
            await self.service.document_repository.save(document)
            
            logger.info(f"项目已从DOCX导入: {project.name}")
            return project
            
        except Exception as e:
            logger.error(f"从DOCX导入项目失败: {e}")
            return None

    async def _do_export_document(self, document: Document, output_path: Path, options: ExportOptions) -> bool:
        """导出文档为DOCX格式"""
        if not DOCX_AVAILABLE:
            logger.error("DOCX库不可用，无法导出DOCX格式")
            return False
            
        try:
            # 创建Word文档
            doc = DocxDocument()
            
            # 添加文档标题
            title = doc.add_heading(document.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加元数据
            if options.include_metadata:
                doc.add_heading('文档信息', level=1)
                info_para = doc.add_paragraph()
                info_para.add_run(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                info_para.add_run(f"文档类型: {document.type}\n")
                info_para.add_run(f"导出工具: AI小说编辑器")
                doc.add_page_break()
            
            # 添加文档内容
            content_paragraphs = document.content.split('\n\n')
            for paragraph_text in content_paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())
            
            # 保存文档
            doc.save(str(output_path))
            
            logger.info(f"文档已导出为DOCX: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"导出文档为DOCX失败: {e}")
            return False

    async def _do_import_document(self, input_path: Path, options: ImportOptions) -> Optional[Document]:
        """从DOCX导入文档"""
        if not DOCX_AVAILABLE:
            logger.error("DOCX库不可用，无法导入DOCX格式")
            return None
            
        try:
            # 验证文件扩展名
            if not self._validate_file_extension(input_path, self.get_supported_extensions()):
                logger.error(f"不支持的文件格式: {input_path.suffix}")
                return None

            # 读取DOCX文件
            doc = DocxDocument(str(input_path))
            
            # 提取标题（使用文件名或第一个标题）
            title = input_path.stem
            
            # 尝试从文档中提取标题
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading') and paragraph.text.strip():
                    title = paragraph.text.strip()
                    break
            
            # 提取所有文本内容
            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # 跳过标题段落
                    if not paragraph.style.name.startswith('Heading') or paragraph.text.strip() != title:
                        content_parts.append(paragraph.text.strip())
            
            content = '\n\n'.join(content_parts)
            
            # 创建文档
            document = Document(
                id="",  # 将由仓储分配
                title=title,
                content=content.strip(),
                type=DocumentType.CHAPTER
            )
            
            logger.info(f"文档已从DOCX导入: {document.title}")
            return document
            
        except Exception as e:
            logger.error(f"从DOCX导入文档失败: {e}")
            return None
